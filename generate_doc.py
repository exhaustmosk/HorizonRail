import sys
import subprocess

# 1. Automatically install python-docx if not present
try:
    import docx
except ImportError:
    print("python-docx not found. Installing it now...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Utility to set background color of a table cell (hex string)"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def create_document():
    doc = Document()

    # Define color palette (Deep Tech Theme)
    COLOR_PRIMARY = RGBColor(124, 58, 237)   # Violet (#7C3AED)
    COLOR_SECONDARY = RGBColor(15, 118, 110) # Teal (#0F766E)
    COLOR_TEXT = RGBColor(30, 41, 59)        # Slate 800 (#1E293B)
    COLOR_MUTED = RGBColor(100, 116, 139)    # Slate 500 (#64748B)

    # Style standard paragraph
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Segoe UI'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXT

    # ─── COVER HEADER ───
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run("HorizonRail")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(4)

    subtitle_p = doc.add_paragraph()
    sub_run = subtitle_p.add_run("Enterprise Alignment & Goal Performance Platform")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = COLOR_SECONDARY
    subtitle_p.paragraph_format.space_after = Pt(24)

    doc.add_paragraph("────────────────────────────────────────────────────────────────").paragraph_format.space_after = Pt(24)

    # ─── SECTION 1: SYSTEM ENDPOINTS & REPO ───
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. System Endpoints & Source Code")
    h1_run.font.size = Pt(16)
    h1_run.font.bold = True
    h1_run.font.color.rgb = COLOR_PRIMARY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)

    # Add quick info table
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    headers = ["Item / Component", "URL Link / Location"]
    data = [
        ("Production Live Deployment", "https://horizonrail.vercel.app"),
        ("Source Code Repository (GitHub)", "https://github.com/exhaustmosk/HorizonRail.git"),
        ("Local Development Server", "http://localhost:5173")
    ]

    # Set column widths
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.3)

    # Populate table with styling
    for idx, (item, val) in enumerate(data):
        row = table.rows[idx]
        
        # Cell 1 (Item)
        c0 = row.cells[0]
        c0.text = item
        c0.paragraphs[0].runs[0].font.bold = True
        c0.paragraphs[0].runs[0].font.color.rgb = COLOR_TEXT
        set_cell_background(c0, "F1F5F9" if idx % 2 == 0 else "FFFFFF")
        
        # Cell 2 (Value)
        c1 = row.cells[1]
        c1.text = val
        c1.paragraphs[0].runs[0].font.color.rgb = COLOR_PRIMARY
        set_cell_background(c1, "F8FAFC" if idx % 2 == 0 else "FFFFFF")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ─── SECTION 2: SYSTEM ARCHITECTURE DIAGRAM ───
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. System Architecture Diagram")
    h2_run.font.size = Pt(16)
    h2_run.font.bold = True
    h2_run.font.color.rgb = COLOR_PRIMARY
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Visual Unicode Flow Diagram representing Architecture
    diagram_p = doc.add_paragraph()
    diagram_p.paragraph_format.space_after = Pt(12)
    
    diagram_text = (
        "┌─────────────────────────────────────────────────────────────────────────┐\n"
        "│                          CLIENT / FRONTEND LAYER                        │\n"
        "│  ┌────────────────────────┐  ┌───────────────────┐  ┌────────────────┐  │\n"
        "│  │    React 18 & Vite     │  │  Zustand Stores   │  │ HTML5 Canvas   │  │\n"
        "│  │ (SPA Routing & Layout) │  │  (Auth & Syncing) │  │ (Neuron Graph) │  │\n"
        "│  └───────────┬────────────┘  └─────────┬─────────┘  └────────┬───────┘  │\n"
        "└──────────────┼─────────────────────────┼─────────────────────┼──────────┘\n"
        "               │                         │                     │           \n"
        "               ▼                         ▼                     ▼           \n"
        "  ⚡ HTTPS / JSON REST API          ⚡ Supabase Web Sockets   ⚡ Real-time JS\n"
        "               │                         │                     │           \n"
        "               └─────────────────────────┼─────────────────────┘           \n"
        "                                         ▼                                 \n"
        "┌─────────────────────────────────────────────────────────────────────────┐\n"
        "│                         SUPABASE BACKEND LAYER                          │\n"
        "│  ┌───────────────────────────────────────────────────────────────────┐  │\n"
        "│  │                     PostgreSQL Database Engine                    │  │\n"
        "│  │  - Organizations (Multi-Tenancy Isolation)                        │  │\n"
        "│  │  - Goal Management (Validation Constraints, State Locks)          │  │\n"
        "│  │  - Audit Logs & Escalation Policy Incidents                       │  │\n"
        "│  └─────────────────────────────────┬─────────────────────────────────┘  │\n"
        "│                                    │ (Postgres Webhook Triggers)        │\n"
        "│                                    ▼                                    │\n"
        "│  ┌───────────────────────────────────────────────────────────────────┐  │\n"
        "│  │                   Supabase Node Edge Functions                    │  │\n"
        "│  │  - Transactional Notification Mailer                              │  │\n"
        "│  │  - Microsoft Teams Webhook Connector                              │  │\n"
        "│  └───────────────────────────────────────────────────────────────────┘  │\n"
        "└─────────────────────────────────────────────────────────────────────────┘"
    )
    
    run_dia = diagram_p.add_run(diagram_text)
    run_dia.font.name = 'Consolas'
    run_dia.font.size = Pt(9.5)
    run_dia.font.color.rgb = COLOR_SECONDARY

    # ─── SECTION 3: ARCHITECTURAL OVERVIEW ───
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Component Breakdown")
    h3_run.font.size = Pt(16)
    h3_run.font.bold = True
    h3_run.font.color.rgb = COLOR_PRIMARY
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(8)

    doc.add_paragraph("The platform uses an advanced decoupled architecture optimized for security, modularity, and rapid real-time state mutations:")

    # Bullet list of key architecture parts
    bullets = [
        ("Multi-Tenant Postgres Engine: ", "Leverages Supabase Schema separation via 'org_id' keys across all critical performance structures, guaranteeing strict security."),
        ("Physics-Based Neuron Alignment Map: ", "A fully-custom React Canvas engine that performs complex physics simulations in 2D space to draw lines of strategic alignment from organizations to Thrust Areas down to individual KPIs."),
        ("Zustand Strict Synced Stores: ", "Maintains client-side synchronization and coordinates role-tab route protection, automatic weight rebalancing calculations, and incident management."),
        ("Compliance & Escalation Policies: ", "An integrated automated system that evaluates goals and check-in schedules, auto-triggering alert incidents when policies are violated."),
        ("Edge Integrations: ", "Asynchronous microservices designed to send webhook notifications to Microsoft Teams channels and transaction notification emails.")
    ]

    for title, desc in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp_run_title = bp.add_run(title)
        bp_run_title.font.bold = True
        bp_run_title.font.color.rgb = COLOR_SECONDARY
        bp.add_run(desc)

    # ─── FOOTER INFO ───
    doc.add_paragraph("────────────────────────────────────────────────────────────────").paragraph_format.space_before = Pt(24)
    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run("HorizonRail Documentation · Generated Automatically")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = COLOR_MUTED

    doc.save("HorizonRail_Architecture.docx")
    print("\n🎉 Success! 'HorizonRail_Architecture.docx' has been successfully created in your workspace root directory!")

if __name__ == "__main__":
    create_document()
